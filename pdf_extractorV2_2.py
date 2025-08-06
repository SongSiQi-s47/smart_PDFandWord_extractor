# -*- coding: utf-8 -*-
"""
PDFå’ŒWordè¡¨æ ¼æå–å·¥å…· - Webç‰ˆæœ¬
ä¿æŒæœ¬åœ°ç‰ˆçš„æ‰€æœ‰åŠŸèƒ½ï¼Œåªå»æ‰input()å‡½æ•°
"""

import os
import re
import logging
import pdfplumber
import pandas as pd
from docx import Document
from typing import List, Dict, Optional, Tuple
from io import BytesIO
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side

# è®¾ç½®æ—¥å¿—
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# æ±‰å­—æ•°å­—
CN_NUM = 'é›¶ä¸€äºŒä¸‰å››äº”å…­ä¸ƒå…«ä¹åç™¾åƒä¸‡äº¿ã€‡å£¹è´°åè‚†ä¼é™†æŸ’æŒç–æ‹¾'

def parse_sample_to_template(sample):
    template = []
    i = 0
    while i < len(sample):
        c = sample[i]
        if c.isdigit():
            num = c
            while i+1 < len(sample) and sample[i+1].isdigit():
                i += 1
                num += sample[i]
            template.append(('digit', num))
        elif c in CN_NUM:
            cn = c
            while i+1 < len(sample) and sample[i+1] in CN_NUM:
                i += 1
                cn += sample[i]
            template.append(('cndigit', cn))
        elif c in '()ï¼ˆï¼‰':
            template.append(('paren', c))
        elif c in 'ã€ã€‘':
            template.append(('bracket', c))
        elif c in '.ã€.':
            template.append(('sep', c))
        elif '\u4e00' <= c <= '\u9fa5':
            word = c
            while i+1 < len(sample) and '\u4e00' <= sample[i+1] <= '\u9fa5':
                i += 1
                word += sample[i]
            template.append(('ch', word))
        else:
            template.append(('other', c))
        i += 1
    return template

def template_to_regex(template):
    regex = ''
    for t, val in template:
        if t == 'digit':
            regex += r'\d+'
        elif t == 'cndigit':
            regex += f'[{CN_NUM}]+'
        elif t == 'paren':
            regex += re.escape(val)
        elif t == 'bracket':
            regex += re.escape(val)
        elif t == 'sep':
            regex += re.escape(val)
        elif t == 'ch':
            regex += f'.{{1,{len(val)+3}}}'
        else:
            regex += re.escape(val)
    return regex

def get_regex_from_sample(sample):
    """ä»æ ·ä¾‹ç›´æ¥è·å–æ­£åˆ™è¡¨è¾¾å¼"""
    template = parse_sample_to_template(sample)
    return template_to_regex(template)

def get_fuzzy_regex_from_sample(sample):
    """
    ç”Ÿæˆæ›´çµæ´»çš„æ­£åˆ™è¡¨è¾¾å¼ï¼Œæ”¯æŒå„ç§ç¼–å·æ ¼å¼
    """
    # åŒ¹é…"æ•°å­—.æ•°å­—.æ•°å­—."ç»“æ„ï¼ˆå¦‚ 9.1.4.3.1.ï¼‰ï¼Œæ›´çµæ´»
    if re.match(r'^\d+(\.\d+)+\.?$', sample):
        # è¿”å›å¸¦åˆ†ç»„çš„æ­£åˆ™ï¼Œç¼–å·éƒ¨åˆ†æ›´çµæ´»ï¼Œæ”¯æŒæœ«å°¾æ²¡æœ‰ç‚¹çš„æƒ…å†µ
        # æ–°å¢ï¼šè®°å½•åŸå§‹æ ·ä¾‹çš„æ•°å­—é•¿åº¦ï¼Œç”¨äºåç»­é•¿åº¦æ£€æŸ¥
        sample_digits = re.sub(r'[^\d]', '', sample)
        regex = re.compile(r'^(\d+(?:\.\d+)+[\.\s\u3000ï¼ã€]*)(.*)$')
        # è¿”å›å­—å…¸ï¼ŒåŒ…å«æ­£åˆ™å’Œé•¿åº¦ä¿¡æ¯
        return {
            'regex': regex,
            'expected_digit_length': len(sample_digits)
        }
    
    # åŒ¹é…"æ•°å­—ï¼‰"æˆ–"æ•°å­—)"ç»“æ„ï¼ˆå¦‚ 1ï¼‰ æˆ– 1) ï¼‰ï¼Œæ”¯æŒä¸­è‹±æ–‡æ‹¬å·
    if re.match(r'^\d+[ï¼‰)]$', sample):
        return {
            'regex': re.compile(r'^(\d+[\)\ï¼‰])(.*)$'),
            'expected_digit_length': None
        }
    
    # åŒ¹é…"ï¼ˆæ•°å­—ï¼‰"æˆ–"(æ•°å­—)"ç»“æ„ï¼ˆå¦‚ ï¼ˆ1ï¼‰ æˆ– (1) ï¼‰ï¼Œæ”¯æŒä¸­è‹±æ–‡æ‹¬å·
    if re.match(r'^[ï¼ˆ(]\d+[ï¼‰)]$', sample):
        return {
            'regex': re.compile(r'^([ï¼ˆ(]\d+[\)\ï¼‰])(.*)$'),
            'expected_digit_length': None
        }
    
    # åŒ¹é…"ï¼ˆæ±‰å­—æ•°å­—ï¼‰"æˆ–"(æ±‰å­—æ•°å­—)"ç»“æ„ï¼ˆå¦‚ ï¼ˆåä¸€ï¼‰ æˆ– (åä¸€) ï¼‰ï¼Œæ”¯æŒæ›´å¤šæ±‰å­—æ•°å­—
    if re.match(r'^[ï¼ˆ(][é›¶ä¸€äºŒä¸‰å››äº”å…­ä¸ƒå…«ä¹åç™¾åƒä¸‡äº¿]+[ï¼‰)]$', sample):
        return {
            'regex': re.compile(r'^([ï¼ˆ(][é›¶ä¸€äºŒä¸‰å››äº”å…­ä¸ƒå…«ä¹åç™¾åƒä¸‡äº¿]+[\)\ï¼‰])(.*)$'),
            'expected_digit_length': None
        }
    
    # æ–°å¢ï¼šåŒ¹é…"æ•°å­—."ç»“æ„ï¼ˆå¦‚ 1.ï¼‰ï¼Œç”Ÿæˆæ›´ä¸¥æ ¼çš„æ­£åˆ™
    if re.match(r'^\d+\.$', sample):
        return {
            'regex': re.compile(r'^(\d+\.)(.*)$'),
            'expected_digit_length': None
        }
    
    # å…¶ä»–ç±»å‹ï¼šå…ˆå»é™¤æ‰€æœ‰ç©ºç™½å­—ç¬¦ï¼ˆåŒ…æ‹¬å…¨è§’ç©ºæ ¼ï¼‰ï¼Œå†ç”¨æ¨¡æ¿è§£æç”Ÿæˆæ­£åˆ™
    sample = re.sub(r'[\s\u3000]', '', sample)
    template = parse_sample_to_template(sample)
    regex = template_to_regex(template)
    regex = regex.rstrip(r'\.')
    regex += r'[\.\s\u3000ï¼ã€]*'  # å…è®¸ç¼–å·åæœ‰ç‚¹ã€ç©ºæ ¼ã€é¡¿å·ç­‰
    return {
        'regex': re.compile(f'^({regex})(.*)$'),
        'expected_digit_length': None
    }

def smart_start_match(sample, text, regex):
    """
    æ™ºèƒ½èµ·å§‹ç¼–å·åŒ¹é…ï¼Œæ”¯æŒå¤šç§åŒ¹é…ç­–ç•¥
    """
    match = regex.match(text)
    if not match:
        return False, None, None
    
    # æå–æ•°å­—åºåˆ—
    sample_digits = re.sub(r'[^\d]', '', sample)
    text_number_part = match.group(1)
    text_digits = re.sub(r'[^\d]', '', text_number_part)
    
    # å¤šç§åŒ¹é…ç­–ç•¥
    if sample_digits == text_digits:
        return (True, "å®Œå…¨åŒ¹é…", text_digits)
    elif text_digits.startswith(sample_digits):
        return (True, "å‰ç¼€åŒ¹é…", text_digits)
    elif sample_digits.startswith(text_digits):
        # å‰ç¼€åŒ¹é…æ—¶æ£€æŸ¥é•¿åº¦æ˜¯å¦ä¸€è‡´
        if len(text_digits) == len(sample_digits):
            return (True, "å‰ç¼€åŒ¹é…", text_digits)
        else:
            return (False, "å‰ç¼€é•¿åº¦ä¸åŒ¹é…", text_digits)
    
    return (False, "ä¸åŒ¹é…", text_digits)

class PDFWordTableExtractor:
    def __init__(self):
        # é»˜è®¤çš„ç›®æ ‡åˆ—æ˜ å°„
        self.target_columns = {
            'åŠŸèƒ½æ¨¡å—': 'ä¸€çº§æ¨¡å—åç§°',
            'åŠŸèƒ½å­é¡¹': 'äºŒçº§æ¨¡å—åç§°',
            'ä¸‰çº§æ¨¡å—': 'ä¸‰çº§æ¨¡å—åç§°',
            'åŠŸèƒ½æè¿°': 'åŠŸèƒ½æè¿°'
        }
        
        # è‡ªå®šä¹‰è¡¨å¤´æ˜ å°„
        self.custom_headers = {}
        
        # çŠ¶æ€å˜é‡
        self.current_lvl1 = ""
        self.current_lvl2 = ""
        self.current_lvl3 = ""
        self.current_description = []
        self.collected_data = []

    def extract_tables(self, file_path: str, file_type: str) -> List[Dict]:
        """æå–è¡¨æ ¼æ•°æ®çš„ä¸»å…¥å£"""
        if file_type == "pdf":
                return self.extract_tables_from_pdf_bid(file_path)
        elif file_type == "docx":
            if "åˆåŒ" in file_path:
                return self.extract_tables_from_word_contract(file_path)
            else:
                return self.extract_tables_from_word_bid(file_path)
        else:
            return []
            
    def setup_custom_headers(self):
        """è®¾ç½®è‡ªå®šä¹‰è¡¨å¤´æ˜ å°„"""
        self.custom_headers = {
            'åŠŸèƒ½æ¨¡å—': 'ä¸€çº§æ¨¡å—åç§°',
            'åŠŸèƒ½å­é¡¹': 'äºŒçº§æ¨¡å—åç§°',
            'ä¸‰çº§æ¨¡å—': 'ä¸‰çº§æ¨¡å—åç§°',
            'åŠŸèƒ½æè¿°': 'åˆåŒæè¿°'
        }
    
    def extract_tables_from_pdf_bid(self, pdf_path: str) -> List[Dict]:
        """æå–PDFæ ‡ä¹¦æ–‡ä»¶ä¸­çš„è¡¨æ ¼ï¼ˆå®Œæ•´ç‰ˆï¼‰"""
        # æ¸…ç©ºä¹‹å‰çš„çŠ¶æ€å˜é‡
        def clear_previous_state():
            self.current_lvl1 = ""
            self.current_lvl2 = ""
            self.current_lvl3 = ""
            self.current_description = []
            self.collected_data = []
        
        clear_previous_state()
        
        def reclassify_module(text, current_level):
            """é‡æ–°åˆ†ç±»æ¨¡å—"""
            if not text:
                return current_level
            
            # æ£€æŸ¥æ˜¯å¦æ˜¯ä¸€çº§æ¨¡å—
            if re.match(r'^[ä¸€äºŒä¸‰å››äº”å…­ä¸ƒå…«ä¹å]+ã€', text) or re.match(r'^\d+\.\d+\.\d+\.\d+', text):
                return 1
            # æ£€æŸ¥æ˜¯å¦æ˜¯äºŒçº§æ¨¡å—
            elif re.match(r'^[ï¼ˆ(][ä¸€äºŒä¸‰å››äº”å…­ä¸ƒå…«ä¹å]+[ï¼‰)]', text) or re.match(r'^\d+\.\d+\.\d+\.\d+\.\d+', text):
                return 2
            # æ£€æŸ¥æ˜¯å¦æ˜¯ä¸‰çº§æ¨¡å—
            elif re.match(r'^\d+[ã€ï¼‰)]', text) or re.match(r'^\d+\.\d+\.\d+\.\d+\.\d+\.\d+', text):
                return 3
            else:
                return current_level
        
        def should_enable_verification():
            """æ˜¯å¦å¯ç”¨éªŒè¯"""
            return True
        
        def is_page_number(text):
            """æ£€æŸ¥æ˜¯å¦ä¸ºé¡µç """
            if not text:
                return False
            page_indicators = ['ç¬¬', 'é¡µ', 'Page', 'page']
            return any(indicator in str(text) for indicator in page_indicators)
        
        data = []
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if not text:
                        continue
                    
                    lines = text.split('\n')
                    current_level = 0
                    
                    for line in lines:
                        line = line.strip()
                        if not line or is_page_number(line):
                            continue
                        
                        # é‡æ–°åˆ†ç±»æ¨¡å—
                        new_level = reclassify_module(line, current_level)
                        if new_level != current_level:
                            # ä¿å­˜ä¹‹å‰çš„æ•°æ®
                            if self.current_description and (self.current_lvl1 or self.current_lvl2 or self.current_lvl3):
                                self.collected_data.append({
                                    'ä¸€çº§æ¨¡å—åç§°': self.current_lvl1,
                                    'äºŒçº§æ¨¡å—åç§°': self.current_lvl2,
                                    'ä¸‰çº§æ¨¡å—åç§°': self.current_lvl3,
                                    'æ ‡ä¹¦æè¿°': ' '.join(self.current_description),
                                    'åˆåŒæè¿°': '',
                                    'æ¥æºæ–‡ä»¶': os.path.basename(pdf_path)
                                })
                            
                            # æ›´æ–°å½“å‰çº§åˆ«
                            current_level = new_level
                            self.current_description = []
                            
                            if current_level == 1:
                                self.current_lvl1 = line
                                self.current_lvl2 = ""
                                self.current_lvl3 = ""
                            elif current_level == 2:
                                self.current_lvl2 = line
                                self.current_lvl3 = ""
                            elif current_level == 3:
                                self.current_lvl3 = line
                        else:
                            # æ”¶é›†æè¿°
                            if current_level > 0:
                                self.current_description.append(line)
                    
                    # å¤„ç†é¡µé¢æœ«å°¾çš„æ•°æ®
                    if self.current_description and (self.current_lvl1 or self.current_lvl2 or self.current_lvl3):
                        self.collected_data.append({
                            'ä¸€çº§æ¨¡å—åç§°': self.current_lvl1,
                            'äºŒçº§æ¨¡å—åç§°': self.current_lvl2,
                            'ä¸‰çº§æ¨¡å—åç§°': self.current_lvl3,
                            'æ ‡ä¹¦æè¿°': ' '.join(self.current_description),
                            'åˆåŒæè¿°': '',
                            'æ¥æºæ–‡ä»¶': os.path.basename(pdf_path)
                        })
                        self.current_description = []
        
        except Exception as e:
            logger.error(f"PDFå¤„ç†é”™è¯¯: {e}")
        
        return self.collected_data

    def extract_tables_from_pdf_bid_with_samples(self, pdf_path: str, lvl1_sample: str, 
                                               lvl2_sample: str = "", lvl3_sample: str = "", 
                                               end_sample: str = "") -> List[Dict]:
        """ä½¿ç”¨ç¼–å·æ ·ä¾‹æå–PDFæ ‡ä¹¦ - ä¸¥æ ¼æŒ‰ç…§æœ¬åœ°ç‰ˆé€»è¾‘"""
        
        # è·å–æ­£åˆ™è¡¨è¾¾å¼
        lvl1_regex_info = get_fuzzy_regex_from_sample(lvl1_sample) if lvl1_sample else None
        lvl2_regex_info = get_fuzzy_regex_from_sample(lvl2_sample) if lvl2_sample else None
        lvl3_regex_info = get_fuzzy_regex_from_sample(lvl3_sample) if lvl3_sample else None
        end_regex_info = get_fuzzy_regex_from_sample(end_sample) if end_sample else None

        lvl1_regex = lvl1_regex_info['regex'] if lvl1_regex_info else None
        lvl2_regex = lvl2_regex_info['regex'] if lvl2_regex_info else None
        lvl3_regex = lvl3_regex_info['regex'] if lvl3_regex_info else None
        end_regex = end_regex_info['regex'] if end_regex_info else None

        results = []
        current_lvl1 = current_lvl2 = current_lvl3 = None
        last_lvl1 = last_lvl2 = last_lvl3 = None
        desc_lines = []
        extracting = False
        start_found = False
        in_lvl3 = False
        in_lvl2 = False

        lvl1_filled = False
        lvl2_filled = False
        lvl1_to_fill = ""
        lvl2_to_fill = ""

        # åˆ¤æ–­æ˜¯å¦æœ‰ä¸‰çº§æ¨¡å—æ ·ä¾‹
        has_lvl3_sample = bool(lvl3_sample)

        def is_page_number(text):
            """åˆ¤æ–­æ˜¯å¦ä¸ºé¡µç ä¿¡æ¯"""
            page_patterns = [
                r'^ç¬¬\d+é¡µ$',
                r'^Page\s*\d+$',
                r'^-\s*\d+\s*-$',
            ]
            # æ›´ä¸¥æ ¼çš„çº¯æ•°å­—é¡µç åˆ¤æ–­
            if re.match(r'^\d+$', text.strip()):
                # å¦‚æœæ•°å­—å°äºç­‰äº3ä½æ•°ï¼Œä¸”å‰åæ²¡æœ‰å…¶ä»–å†…å®¹ï¼Œå¯èƒ½æ˜¯é¡µç 
                if len(text.strip()) <= 3:
                    return True
            return any(re.match(pattern, text.strip()) for pattern in page_patterns)

        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    lines = (page.extract_text() or '').split('\n')
                    for raw_text in lines:
                        # å¤„ç†åŸå§‹æ–‡æœ¬
                        text = re.sub(r'[\s\u3000]', '', raw_text)

                        # ç»ˆæ­¢ç¼–å·åˆ¤æ–­
                        if end_regex and end_regex.match(text):
                            is_match, match_type, actual_digits = smart_start_match(end_sample, text, end_regex)
                            if is_match:
                                extracting = False
                                in_lvl3 = False
                                in_lvl2 = False
                                continue

                        # æ™ºèƒ½èµ·å§‹ç¼–å·åˆ¤æ–­
                        if not extracting and lvl1_sample and lvl1_regex and lvl1_regex.match(text):
                            is_match, match_type, actual_digits = smart_start_match(lvl1_sample, text, lvl1_regex)
                            if is_match:
                                extracting = True
                                start_found = True
                                current_lvl1 = raw_text.strip()
                                lvl1_filled = False
                                lvl1_to_fill = current_lvl1
                                lvl2_to_fill = current_lvl2 if current_lvl2 else ""
                                in_lvl2 = False
                                continue

                        if not extracting:
                            continue

                        # å…ˆåˆ¤æ–­ä¸‰çº§
                        m3 = lvl3_regex.match(text) if lvl3_regex else None
                        if m3:
                            # éªŒè¯ä¸‰çº§æ¨¡å—åŒ¹é…çš„æœ‰æ•ˆæ€§
                            def is_valid_lvl3_match(match_obj, original_text):
                                """éªŒè¯ä¸‰çº§æ¨¡å—åŒ¹é…æ˜¯å¦æœ‰æ•ˆ"""
                                if not match_obj:
                                    return False
                                
                                # æå–åŒ¹é…çš„ç¼–å·éƒ¨åˆ†
                                number_part = match_obj.group(1)
                                title_part = match_obj.group(2).strip()
                                
                                # æ£€æŸ¥æ ‡é¢˜éƒ¨åˆ†æ˜¯å¦æœ‰å†…å®¹
                                if not title_part:
                                    return False
                                
                                return True
                            
                            # éªŒè¯åŒ¹é…æœ‰æ•ˆæ€§
                            if not is_valid_lvl3_match(m3, raw_text.strip()):
                                m3 = None
                            
                            if lvl3_regex_info and lvl3_regex_info['expected_digit_length']:
                                text_digits = re.sub(r'[^\d]', '', m3.group(1))
                                if len(text_digits) != lvl3_regex_info['expected_digit_length']:
                                    m3 = None

                            if m3:
                                # é‡åˆ°æ–°ä¸‰çº§ç¼–å·æ—¶ï¼Œå…ˆè¾“å‡ºä¸Šä¸€ç»„ï¼ˆå¦‚æœæœ‰æè¿°ï¼‰
                                if in_lvl3 and desc_lines:
                                    results.append({
                                        "ä¸€çº§æ¨¡å—åç§°": lvl1_to_fill,
                                        "äºŒçº§æ¨¡å—åç§°": lvl2_to_fill,
                                        "ä¸‰çº§æ¨¡å—åç§°": last_lvl3,
                                        "æ ‡ä¹¦æè¿°": '\n\n'.join(self._merge_paragraphs(desc_lines)).strip(),
                                        "åˆåŒæè¿°": "",
                                        "æ¥æºæ–‡ä»¶": os.path.basename(pdf_path)
                                    })
                                    desc_lines = []
                                    lvl1_filled = True
                                    lvl2_filled = True

                                number = m3.group(1)
                                title = m3.group(2).strip()
                                current_lvl3 = f"{number} {title}".strip()
                                last_lvl3 = current_lvl3
                                in_lvl3 = True
                                in_lvl2 = False
                                lvl1_to_fill = current_lvl1 if not lvl1_filled else ""
                                lvl2_to_fill = current_lvl2 if not lvl2_filled else ""
                                continue

                        # å†åˆ¤æ–­äºŒçº§
                        m2 = lvl2_regex.match(text) if lvl2_regex else None
                        if m2:
                            # éªŒè¯äºŒçº§æ¨¡å—åŒ¹é…çš„æœ‰æ•ˆæ€§
                            def is_valid_lvl2_match(match_obj, original_text):
                                """éªŒè¯äºŒçº§æ¨¡å—åŒ¹é…æ˜¯å¦æœ‰æ•ˆ"""
                                if not match_obj:
                                    return False
                                
                                # æå–åŒ¹é…çš„ç¼–å·éƒ¨åˆ†
                                number_part = match_obj.group(1)
                                title_part = match_obj.group(2).strip()
                                
                                # æ£€æŸ¥æ ‡é¢˜éƒ¨åˆ†æ˜¯å¦æœ‰å†…å®¹
                                if not title_part:
                                    return False
                                
                                return True
                            
                            # éªŒè¯äºŒçº§æ¨¡å—åŒ¹é…æœ‰æ•ˆæ€§
                            if not is_valid_lvl2_match(m2, raw_text.strip()):
                                m2 = None
                            
                            if lvl2_regex_info and lvl2_regex_info['expected_digit_length']:
                                text_digits = re.sub(r'[^\d]', '', m2.group(1))
                                if len(text_digits) != lvl2_regex_info['expected_digit_length']:
                                    m2 = None
                            
                            if m2:
                                # æ ¹æ®æ˜¯å¦æœ‰ä¸‰çº§æ¨¡å—ä½¿ç”¨ä¸åŒé€»è¾‘
                                if has_lvl3_sample:
                                    # æœ‰ä¸‰çº§æ¨¡å—æ ·ä¾‹ï¼šä½¿ç”¨è€ä»£ç é€»è¾‘
                                    if in_lvl3 and desc_lines:
                                        results.append({
                                            "ä¸€çº§æ¨¡å—åç§°": lvl1_to_fill,
                                            "äºŒçº§æ¨¡å—åç§°": lvl2_to_fill,
                                            "ä¸‰çº§æ¨¡å—åç§°": last_lvl3,
                                            "æ ‡ä¹¦æè¿°": '\n\n'.join(self._merge_paragraphs(desc_lines)).strip(),
                                            "åˆåŒæè¿°": "",
                                            "æ¥æºæ–‡ä»¶": os.path.basename(pdf_path)
                                        })
                                        desc_lines = []
                                        lvl1_filled = True
                                        lvl2_filled = True
                                        in_lvl3 = False

                                    number = m2.group(1)
                                    title = m2.group(2).strip()
                                    current_lvl2 = f"{number} {title}".strip()
                                    current_lvl3 = None
                                    last_lvl2 = current_lvl2
                                    lvl2_filled = False
                                    in_lvl2 = True
                                    in_lvl3 = False
                                    
                                    # æ›´æ–°å¡«å……å€¼
                                    lvl1_to_fill = current_lvl1 if not lvl1_filled else ""
                                    lvl2_to_fill = current_lvl2
                                    continue
                                else:
                                    # æ²¡æœ‰ä¸‰çº§æ¨¡å—æ ·ä¾‹ï¼šä½¿ç”¨æ–°ä»£ç é€»è¾‘
                                    # å¦‚æœä¹‹å‰æœ‰æè¿°å†…å®¹ï¼Œå…ˆè¾“å‡ºä¸Šä¸€ç»„
                                    if desc_lines and in_lvl2 and last_lvl2:
                                        results.append({
                                            "ä¸€çº§æ¨¡å—åç§°": lvl1_to_fill,
                                            "äºŒçº§æ¨¡å—åç§°": last_lvl2,  # ä½¿ç”¨ä¸Šä¸€ä¸ªäºŒçº§æ¨¡å—åç§°
                                            "ä¸‰çº§æ¨¡å—åç§°": "",
                                            "æ ‡ä¹¦æè¿°": '\n\n'.join(self._merge_paragraphs(desc_lines)).strip(),
                                            "åˆåŒæè¿°": "",
                                            "æ¥æºæ–‡ä»¶": os.path.basename(pdf_path)
                                        })
                                        desc_lines = []
                                        lvl1_filled = True
                                        lvl2_filled = True
                                    
                                    # æ›´æ–°å½“å‰äºŒçº§æ¨¡å—
                                    number = m2.group(1)
                                    title = m2.group(2).strip()
                                    current_lvl2 = f"{number} {title}".strip()
                                    current_lvl3 = None
                                    last_lvl2 = current_lvl2
                                    lvl2_filled = False
                                    in_lvl2 = True
                                    in_lvl3 = False
                                    
                                    # æ›´æ–°å¡«å……å€¼
                                    lvl1_to_fill = current_lvl1 if not lvl1_filled else ""
                                    lvl2_to_fill = current_lvl2
                                    continue

                        # æœ€ååˆ¤æ–­ä¸€çº§
                        m1 = lvl1_regex.match(text) if lvl1_regex else None
                        if m1:
                            # éªŒè¯ä¸€çº§æ¨¡å—åŒ¹é…çš„æœ‰æ•ˆæ€§
                            def is_valid_lvl1_match(match_obj, original_text):
                                """éªŒè¯ä¸€çº§æ¨¡å—åŒ¹é…æ˜¯å¦æœ‰æ•ˆ"""
                                if not match_obj:
                                    return False
                                
                                # æå–åŒ¹é…çš„ç¼–å·éƒ¨åˆ†
                                number_part = match_obj.group(1)
                                title_part = match_obj.group(2).strip()
                                
                                # æ£€æŸ¥æ ‡é¢˜éƒ¨åˆ†æ˜¯å¦æœ‰å†…å®¹
                                if not title_part:
                                    return False
                                
                                return True
                            
                            # éªŒè¯ä¸€çº§æ¨¡å—åŒ¹é…æœ‰æ•ˆæ€§
                            if not is_valid_lvl1_match(m1, raw_text.strip()):
                                m1 = None
                            
                            if lvl1_regex_info and lvl1_regex_info['expected_digit_length']:
                                text_digits = re.sub(r'[^\d]', '', m1.group(1))
                                if len(text_digits) != lvl1_regex_info['expected_digit_length']:
                                    m1 = None
                            
                            if m1:
                                # æ ¹æ®æ˜¯å¦æœ‰ä¸‰çº§æ¨¡å—ä½¿ç”¨ä¸åŒé€»è¾‘
                                if has_lvl3_sample:
                                    # æœ‰ä¸‰çº§æ¨¡å—æ ·ä¾‹ï¼šä½¿ç”¨è€ä»£ç é€»è¾‘
                                    if in_lvl3 and desc_lines:
                                        results.append({
                                            "ä¸€çº§æ¨¡å—åç§°": lvl1_to_fill,
                                            "äºŒçº§æ¨¡å—åç§°": lvl2_to_fill,
                                            "ä¸‰çº§æ¨¡å—åç§°": last_lvl3,
                                            "æ ‡ä¹¦æè¿°": '\n\n'.join(self._merge_paragraphs(desc_lines)).strip(),
                                            "åˆåŒæè¿°": "",
                                            "æ¥æºæ–‡ä»¶": os.path.basename(pdf_path)
                                        })
                                        desc_lines = []
                                        lvl1_filled = True
                                        lvl2_filled = True
                                        in_lvl3 = False
                                else:
                                    # æ²¡æœ‰ä¸‰çº§æ¨¡å—æ ·ä¾‹ï¼šä½¿ç”¨æ–°ä»£ç é€»è¾‘
                                    if desc_lines and not in_lvl3:
                                        results.append({
                                            "ä¸€çº§æ¨¡å—åç§°": lvl1_to_fill,
                                            "äºŒçº§æ¨¡å—åç§°": lvl2_to_fill,
                                            "ä¸‰çº§æ¨¡å—åç§°": "",
                                            "æ ‡ä¹¦æè¿°": '\n\n'.join(self._merge_paragraphs(desc_lines)).strip(),
                                            "åˆåŒæè¿°": "",
                                            "æ¥æºæ–‡ä»¶": os.path.basename(pdf_path)
                                        })
                                        desc_lines = []
                                        lvl1_filled = True
                                        lvl2_filled = True

                                    if in_lvl3 and desc_lines:
                                        results.append({
                                            "ä¸€çº§æ¨¡å—åç§°": lvl1_to_fill,
                                            "äºŒçº§æ¨¡å—åç§°": lvl2_to_fill,
                                            "ä¸‰çº§æ¨¡å—åç§°": last_lvl3,
                                            "æ ‡ä¹¦æè¿°": '\n\n'.join(self._merge_paragraphs(desc_lines)).strip(),
                                            "åˆåŒæè¿°": "",
                                            "æ¥æºæ–‡ä»¶": os.path.basename(pdf_path)
                                        })
                                        desc_lines = []
                                        lvl1_filled = True
                                        lvl2_filled = True
                                        in_lvl3 = False

                                number = m1.group(1)
                                title = m1.group(2).strip()
                                current_lvl1 = f"{number} {title}".strip()
                                current_lvl2 = current_lvl3 = None
                                last_lvl1 = current_lvl1
                                lvl1_filled = False
                                lvl2_filled = False
                                in_lvl2 = False
                                in_lvl3 = False
                                lvl1_to_fill = current_lvl1
                                lvl2_to_fill = current_lvl2 if current_lvl2 else ""
                                continue

                        # ä¿®æ”¹åçš„æ”¶é›†é€»è¾‘ - ä¸¥æ ¼æŒ‰ç…§æœ¬åœ°ç‰ˆ
                        if extracting:
                            # æ›´ä¸¥æ ¼çš„æè¿°æ”¶é›†éªŒè¯
                            def should_collect_description():
                                """åˆ¤æ–­æ˜¯å¦åº”è¯¥æ”¶é›†æè¿°å†…å®¹"""
                                # å¦‚æœæœ‰ä¸‰çº§æ¨¡å—æ ·ä¾‹
                                if has_lvl3_sample:
                                    # åªæœ‰åœ¨ä¸‰çº§æ¨¡å—ä¸‹æ‰æ”¶é›†
                                    return in_lvl3
                                else:
                                    # æ²¡æœ‰ä¸‰çº§æ¨¡å—æ ·ä¾‹æ—¶
                                    if in_lvl2:
                                        return True
                                    return False
                            
                            # åªæœ‰åœ¨åº”è¯¥æ”¶é›†çš„æƒ…å†µä¸‹æ‰æ·»åŠ æè¿°
                            if should_collect_description():
                                # é¢å¤–éªŒè¯ï¼šç¡®ä¿ä¸æ˜¯æ¨¡å—æ ‡é¢˜è¡Œ
                                if not (lvl1_regex and lvl1_regex.match(text)) and \
                                   not (lvl2_regex and lvl2_regex.match(text)) and \
                                   not (lvl3_regex and lvl3_regex.match(text)):
                                    desc_lines.append(raw_text.strip())

            # è¡¥å……æœ€åä¸€ç»„
            if in_lvl3 and desc_lines:
                results.append({
                    "ä¸€çº§æ¨¡å—åç§°": lvl1_to_fill,
                    "äºŒçº§æ¨¡å—åç§°": lvl2_to_fill,
                    "ä¸‰çº§æ¨¡å—åç§°": last_lvl3,
                    "æ ‡ä¹¦æè¿°": '\n\n'.join(self._merge_paragraphs(desc_lines)).strip(),
                    "åˆåŒæè¿°": "",
                    "æ¥æºæ–‡ä»¶": os.path.basename(pdf_path)
                })
            elif desc_lines:  # å¦‚æœæ²¡æœ‰ä¸‰çº§æ¨¡å—ä½†æœ‰æè¿°å†…å®¹
                # ç¡®ä¿æœ‰æ­£ç¡®çš„äºŒçº§æ¨¡å—åç§°
                final_lvl2 = last_lvl2 if last_lvl2 else current_lvl2 if current_lvl2 else lvl2_to_fill
                results.append({
                    "ä¸€çº§æ¨¡å—åç§°": current_lvl1 if current_lvl1 else lvl1_to_fill,
                    "äºŒçº§æ¨¡å—åç§°": final_lvl2,
                    "ä¸‰çº§æ¨¡å—åç§°": "",
                    "æ ‡ä¹¦æè¿°": '\n\n'.join(self._merge_paragraphs(desc_lines)).strip(),
                    "åˆåŒæè¿°": "",
                    "æ¥æºæ–‡ä»¶": os.path.basename(pdf_path)
                })

            # æ¸…ç†æ•°æ®
            results = [r for r in results if any([r["ä¸€çº§æ¨¡å—åç§°"], r["äºŒçº§æ¨¡å—åç§°"], r["ä¸‰çº§æ¨¡å—åç§°"], r["æ ‡ä¹¦æè¿°"]])]
            
        except Exception as e:
            logger.error(f"PDFå¤„ç†é”™è¯¯: {e}")
        
        return results

    def extract_tables_from_pdf_contract_with_samples(self, pdf_path: str, lvl1_sample: str, 
                                                    lvl2_sample: str = "", lvl3_sample: str = "", 
                                                    end_sample: str = "") -> List[Dict]:
        """ä½¿ç”¨ç¼–å·æ ·ä¾‹æå–PDFåˆåŒ"""
        data = self.extract_tables_from_pdf_bid_with_samples(pdf_path, lvl1_sample, lvl2_sample, lvl3_sample, end_sample)
        
        # è°ƒæ•´æè¿°å­—æ®µæ˜ å°„
        for item in data:
            if item['æ ‡ä¹¦æè¿°']:
                item['åˆåŒæè¿°'] = item['æ ‡ä¹¦æè¿°']
                item['æ ‡ä¹¦æè¿°'] = ''
        
        return data
    
    def extract_tables_from_word_contract(self, docx_path: str, original_filename: str = None) -> List[Dict]:
        """æå–WordåˆåŒæ–‡ä»¶ä¸­çš„è¡¨æ ¼ï¼ˆå®Œæ•´ç‰ˆï¼‰"""
        data = []
        found_quotation_section = False
        current_headers = None
        doc = Document(docx_path)
        
        # å¦‚æœæ²¡æœ‰æä¾›åŸå§‹æ–‡ä»¶åï¼Œä½¿ç”¨docx_path
        if not original_filename:
            original_filename = os.path.basename(docx_path)
        
        # å…ˆæ£€æŸ¥æ•´ä¸ªæ–‡æ¡£æ˜¯å¦åŒ…å«åˆ†é¡¹æŠ¥ä»·è¡¨
        has_quotation_table = False
        for para in doc.paragraphs:
            if "åˆ†é¡¹æŠ¥ä»·è¡¨" in para.text:
                has_quotation_table = True
                break
        
        if not has_quotation_table:
            return data
        
        # å¤„ç†æ‰€æœ‰è¡¨æ ¼
        for table_idx, table in enumerate(doc.tables):
            rows = list(table.rows)
            if not rows:
                continue
            
            # æ£€æŸ¥è¡¨å¤´
            headers = [cell.text.strip().replace('\n', '') for cell in rows[0].cells]
            
            # æ£€æŸ¥æ˜¯å¦æ˜¯ç›®æ ‡è¡¨æ ¼
            if self._is_target_table_custom(headers):
                current_headers = headers
                found_quotation_section = True
                start_row = 1
                print(f"âœ… æ‰¾åˆ°åŒ¹é…çš„è¡¨æ ¼ï¼Œè¡¨å¤´ï¼š{headers}")
            elif current_headers and found_quotation_section:
                start_row = 0
            else:
                continue
                
            # å¤„ç†æ•°æ®è¡Œ
            for row_idx, row in enumerate(rows[start_row:], start=start_row):
                row_data = {}
                cells = row.cells
                
                # å¤„ç†åˆå¹¶å•å…ƒæ ¼çš„æƒ…å†µ
                for idx, header in enumerate(current_headers):
                    if idx < len(cells):
                        cell_text = cells[idx].text.strip()
                        row_data[header] = cell_text
                    else:
                        row_data[header] = ''
                
                # æ£€æŸ¥æ˜¯å¦æœ‰æœ‰æ•ˆæ•°æ®
                has_data = False
                
                # æ£€æŸ¥åºå·å­—æ®µ
                for header in current_headers:
                    if 'åºå·' in header and row_data.get(header, '').strip():
                        try:
                            int(row_data[header])
                            has_data = True
                            break
                        except ValueError:
                            pass
                
                # å¦‚æœåºå·ä¸æ˜¯æ•°å­—ï¼Œæ£€æŸ¥å…¶ä»–å…³é”®å­—æ®µ
                if not has_data:
                    key_fields = self._get_key_fields_for_check()
                    for field in key_fields:
                        for header in current_headers:
                            if field in header and row_data.get(header, '').strip():
                                has_data = True
                                break
                        if has_data:
                            break
                
                # å¦‚æœå…³é”®å­—æ®µéƒ½æ²¡æœ‰ï¼Œå†æ£€æŸ¥å…¶ä»–å­—æ®µ
                if not has_data:
                    has_data = any(row_data.get(header, '').strip() for header in current_headers)
                
                # æ·»åŠ è°ƒè¯•ä¿¡æ¯
                if has_data:
                    mapped = self._map_word_row_custom(row_data, docx_path, original_filename)
                    
                    # æ£€æŸ¥é‡å¤å¹¶å¤„ç†
                    if len(data) > 0:
                        # æŸ¥æ‰¾ä¸Šä¸€ä¸ªéç©ºçš„ä¸€çº§æ¨¡å—åç§°
                        last_lvl1 = ""
                        last_lvl2 = ""
                        for i in range(len(data) - 1, -1, -1):
                            if data[i]['ä¸€çº§æ¨¡å—åç§°'].strip():
                                last_lvl1 = data[i]['ä¸€çº§æ¨¡å—åç§°']
                                break
                        for i in range(len(data) - 1, -1, -1):
                            if data[i]['äºŒçº§æ¨¡å—åç§°'].strip():
                                last_lvl2 = data[i]['äºŒçº§æ¨¡å—åç§°']
                                break
                        
                        # æ£€æŸ¥ä¸€çº§æ¨¡å—åç§°æ˜¯å¦é‡å¤
                        if mapped['ä¸€çº§æ¨¡å—åç§°'].strip() == last_lvl1.strip() and mapped['ä¸€çº§æ¨¡å—åç§°'].strip():
                            # å¦‚æœä¸€çº§æ¨¡å—åç§°é‡å¤ï¼Œæ¸…ç©ºå½“å‰çš„ä¸€çº§æ¨¡å—åç§°
                            mapped['ä¸€çº§æ¨¡å—åç§°'] = ""
                        
                        # æ£€æŸ¥äºŒçº§æ¨¡å—åç§°æ˜¯å¦é‡å¤
                        if mapped['äºŒçº§æ¨¡å—åç§°'].strip() == last_lvl2.strip() and mapped['äºŒçº§æ¨¡å—åç§°'].strip():
                            # å¦‚æœäºŒçº§æ¨¡å—åç§°é‡å¤ï¼Œæ¸…ç©ºå½“å‰çš„äºŒçº§æ¨¡å—åç§°
                            mapped['äºŒçº§æ¨¡å—åç§°'] = ""
                    
                    data.append(mapped)
        
        return data

    def extract_tables_from_word_bid(self, docx_path: str) -> List[Dict]:
        """æå–Wordæ ‡ä¹¦æ–‡ä»¶ä¸­çš„è¡¨æ ¼"""
        data = self.extract_tables_from_word_contract(docx_path)
        
        # è°ƒæ•´æè¿°å­—æ®µæ˜ å°„
        for item in data:
            if item['åˆåŒæè¿°']:
                item['æ ‡ä¹¦æè¿°'] = item['åˆåŒæè¿°']
                item['åˆåŒæè¿°'] = ''
        
        return data
        
    def _is_target_table_custom(self, headers: List[str]) -> bool:
        """æ£€æŸ¥æ˜¯å¦æ˜¯ç›®æ ‡è¡¨æ ¼ï¼ˆè‡ªå®šä¹‰è¡¨å¤´ï¼‰"""
        if self.custom_headers:
            # ä½¿ç”¨è‡ªå®šä¹‰è¡¨å¤´
            found_headers = 0
            for custom_header in self.custom_headers.keys():
                for header in headers:
                    if custom_header in header:
                        found_headers += 1
                        break
            return found_headers >= 2
        else:
            # ä½¿ç”¨é»˜è®¤è¡¨å¤´
            return self._is_target_table(headers)
    
    def _get_key_fields_for_check(self) -> List[str]:
        """è·å–ç”¨äºæ£€æŸ¥çš„å…³é”®å­—æ®µåˆ—è¡¨"""
        if self.custom_headers:
            # ä½¿ç”¨è‡ªå®šä¹‰è¡¨å¤´
            return list(self.custom_headers.keys())
        else:
            # ä½¿ç”¨é»˜è®¤è¡¨å¤´
            return ['åŠŸèƒ½æè¿°', 'ä¸‰çº§æ¨¡å—', 'åŠŸèƒ½æ¨¡å—', 'åŠŸèƒ½å­é¡¹']
    
    def _map_word_row_custom(self, row_data: Dict, source_file: str, original_filename: str = None) -> Dict:
        """ä½¿ç”¨è‡ªå®šä¹‰è¡¨å¤´æ˜ å°„Wordè¡Œæ•°æ®"""
        if self.custom_headers:
            # ä½¿ç”¨è‡ªå®šä¹‰è¡¨å¤´æ˜ å°„
            mapped_data = {}
            for header, value in row_data.items():
                for custom_header, standard_field in self.custom_headers.items():
                    if custom_header in header:
                        mapped_data[standard_field] = value
                        break
            
            # ç¡®ä¿WordåˆåŒå†…å®¹æ˜ å°„åˆ°"åˆåŒæè¿°"
            desc_value = mapped_data.get('åˆåŒæè¿°', '')
            
            mapped = {
                'ä¸€çº§æ¨¡å—åç§°': mapped_data.get('ä¸€çº§æ¨¡å—åç§°', ''),
                'äºŒçº§æ¨¡å—åç§°': mapped_data.get('äºŒçº§æ¨¡å—åç§°', ''),
                'ä¸‰çº§æ¨¡å—åç§°': mapped_data.get('ä¸‰çº§æ¨¡å—åç§°', ''),
                'æ ‡ä¹¦æè¿°': '',  # WordåˆåŒæ–‡ä»¶ï¼Œæ ‡ä¹¦æè¿°ä¸ºç©º
                'åˆåŒæè¿°': desc_value,  # WordåˆåŒå†…å®¹æ”¾è¿™é‡Œ
                'æ¥æºæ–‡ä»¶': original_filename if original_filename else os.path.basename(source_file),
            }
        else:
            # ä½¿ç”¨é»˜è®¤æ˜ å°„
            mapped = self._map_word_row(row_data, source_file)
        
        return mapped
    
    def _is_target_table(self, headers: List[str]) -> bool:
        """æ£€æŸ¥æ˜¯å¦æ˜¯ç›®æ ‡è¡¨æ ¼ï¼ˆé»˜è®¤è¡¨å¤´ï¼‰"""
        target_headers = list(self.target_columns.keys())
        found_headers = []
        
        # é¢„å¤„ç†ç›®æ ‡å­—æ®µåï¼Œå»æ‰æ¢è¡Œç¬¦
        normalized_target_headers = {}
        for target in target_headers:
            normalized_target = target.replace('\n', '').replace('\\n', '')
            normalized_target_headers[normalized_target] = target
        
        for normalized_target, original_target in normalized_target_headers.items():
            for header_cell in headers:
                # å¯¹å®é™…å­—æ®µåä¹Ÿè¿›è¡Œé¢„å¤„ç†
                normalized_header = header_cell.replace('\n', '').replace('\\n', '')
                if normalized_target in normalized_header:
                    found_headers.append(original_target)
                    break
        
        return len(found_headers) >= 2
    
    def _map_word_row(self, row_data: Dict, source_file: str) -> Dict:
        """æ˜ å°„Wordè¡Œæ•°æ®åˆ°æ ‡å‡†æ ¼å¼"""
        # åˆ›å»ºæ ‡å‡†åŒ–å­—æ®µæ˜ å°„
        field_mapping = {}
        for original_field in ['åŠŸèƒ½æ¨¡å—', 'åŠŸèƒ½å­é¡¹', 'ä¸‰çº§æ¨¡å—', 'åŠŸèƒ½æè¿°']:
            normalized_field = original_field.replace('\n', '').replace('\\n', '')
            field_mapping[normalized_field] = original_field
        
        # æŸ¥æ‰¾åŒ¹é…çš„å­—æ®µ
        mapped_data = {}
        for header, value in row_data.items():
            normalized_header = header.replace('\n', '').replace('\\n', '')
            for normalized_field, original_field in field_mapping.items():
                if normalized_field in normalized_header:
                    mapped_data[original_field] = value
                    break
        
        mapped = {
            'ä¸€çº§æ¨¡å—åç§°': mapped_data.get('åŠŸèƒ½æ¨¡å—', ''),
            'äºŒçº§æ¨¡å—åç§°': mapped_data.get('åŠŸèƒ½å­é¡¹', ''),
            'ä¸‰çº§æ¨¡å—åç§°': mapped_data.get('ä¸‰çº§æ¨¡å—', ''),
            'æ ‡ä¹¦æè¿°': '',
            'åˆåŒæè¿°': '',
            'æ¥æºæ–‡ä»¶': os.path.basename(source_file) if not source_file.endswith('tmp') else 'åˆåŒ.docx',
        }
        
        # æ ¹æ®æ–‡ä»¶ç±»å‹å†³å®šå†…å®¹æ”¾å“ªä¸ªå­—æ®µ
        if "æ ‡ä¹¦" in source_file:
            mapped['æ ‡ä¹¦æè¿°'] = mapped_data.get('åŠŸèƒ½æè¿°', '')
        elif "åˆåŒ" in source_file:
            mapped['åˆåŒæè¿°'] = mapped_data.get('åŠŸèƒ½æè¿°', '')
        
        return mapped
    
    def _merge_paragraphs(self, desc_lines):
        """åˆå¹¶è‡ªç„¶æ®µ"""
        desc_paragraphs = []
        current_para = []
        for line in desc_lines:
            if not line.strip():
                if current_para:
                    desc_paragraphs.append(''.join(current_para))
                    current_para = []
            else:
                current_para.append(line.strip())
        if current_para:
            desc_paragraphs.append(''.join(current_para))
        return desc_paragraphs
    
    def _clean_extracted_data(self, data: List[Dict]) -> List[Dict]:
        """æ¸…ç†æå–çš„æ•°æ®ï¼Œè¿‡æ»¤æ— æ•ˆå†…å®¹"""
        cleaned_data = []
        
        for item in data:
            # æ£€æŸ¥æ˜¯å¦ä¸ºé¡µç ä¿¡æ¯
            def is_page_content(text):
                if not text:
                    return False
                page_indicators = ['ç¬¬', 'é¡µ', 'Page', 'page']
                return any(indicator in str(text) for indicator in page_indicators)
            
            # æ¸…ç†å„å­—æ®µ
            cleaned_item = {}
            for key, value in item.items():
                if is_page_content(value):
                    print(f"DEBUG: è¿‡æ»¤é¡µç å†…å®¹ '{value}' ä»å­—æ®µ '{key}'")
                    cleaned_item[key] = ""
                else:
                    cleaned_item[key] = value
            
            # åªä¿ç•™æœ‰æ•ˆæ•°æ®
            if any([cleaned_item.get("ä¸€çº§æ¨¡å—åç§°"), cleaned_item.get("äºŒçº§æ¨¡å—åç§°"), 
                   cleaned_item.get("ä¸‰çº§æ¨¡å—åç§°"), cleaned_item.get("æ ‡ä¹¦æè¿°"), cleaned_item.get("åˆåŒæè¿°")]):
                cleaned_data.append(cleaned_item)
        
        return cleaned_data
    
    def create_excel_output(self, data: List[Dict], output_path: str, append_mode=False):
        """åˆ›å»ºExcelè¾“å‡ºæ–‡ä»¶"""
        if not data:
            logger.warning("æ²¡æœ‰æ•°æ®éœ€è¦è¾“å‡º")
            return None
        
        # æ–°å¢ï¼šæ™ºèƒ½åˆ†æ®µå¤„ç†å‡½æ•°
        def split_long_description(description, max_length=500):
            """æ™ºèƒ½åˆ†æ®µå¤„ç†é•¿æè¿°"""
            if not description or len(description) <= max_length:
                return [description]
            
            # é¢„å®šä¹‰ç¼–å·æ ¼å¼æ­£åˆ™è¡¨è¾¾å¼
            number_patterns = [
                # ä¸­æ–‡æ•°å­—æ ¼å¼
                r'[ä¸€äºŒä¸‰å››äº”å…­ä¸ƒå…«ä¹å]+ã€',  # ä¸€ã€äºŒã€ä¸‰ã€
                r'[ï¼ˆ(][ä¸€äºŒä¸‰å››äº”å…­ä¸ƒå…«ä¹å]+[ï¼‰)]',  # ï¼ˆä¸€ï¼‰ï¼ˆäºŒï¼‰
                r'[ï¼ˆ(]\d+[ï¼‰)]',  # ï¼ˆ1ï¼‰ï¼ˆ2ï¼‰
                r'\d+ã€',  # 1ã€2ã€
                r'\d+\.',  # 1.2.3.
                r'\d+ï¼‰',  # 1ï¼‰2ï¼‰
                r'\d+\)',  # 1)2)
                
                # è‹±æ–‡æ•°å­—æ ¼å¼
                r'\d+\.',  # 1.2.3.
                r'[a-z]+\)',  # a)b)c)
                r'[A-Z]+\.',  # A.B.C.
                r'[i]+\)',  # i)ii)iii)
            ]
            
            # å°è¯•æŒ‰ç¼–å·åˆ†å‰²
            for pattern in number_patterns:
                parts = re.split(f'({pattern})', description)
                if len(parts) > 1:
                    # é‡æ–°ç»„åˆåˆ†å‰²çš„éƒ¨åˆ†
                    result = []
                    current_part = ""
                    for i, part in enumerate(parts):
                        if re.match(pattern, part):
                            if current_part:
                                result.append(current_part.strip())
                            current_part = part
                        else:
                            current_part += part
                    
                    if current_part:
                        result.append(current_part.strip())
                    
                    # å¦‚æœåˆ†å‰²åçš„éƒ¨åˆ†ä»ç„¶å¤ªé•¿ï¼Œç»§ç»­åˆ†å‰²
                    final_result = []
                    for part in result:
                        if len(part) > max_length:
                            # æŒ‰å¥å·åˆ†å‰²
                            sentences = re.split(r'[ã€‚ï¼ï¼Ÿ]', part)
                            current_sentence = ""
                            for sentence in sentences:
                                if len(current_sentence + sentence) <= max_length:
                                    current_sentence += sentence + "ã€‚"
                                else:
                                    if current_sentence:
                                        final_result.append(current_sentence.strip())
                                    current_sentence = sentence + "ã€‚"
                            if current_sentence:
                                final_result.append(current_sentence.strip())
                        else:
                            final_result.append(part)
                    
                    return final_result
            
            # å¦‚æœæ²¡æœ‰æ‰¾åˆ°ç¼–å·æ ¼å¼ï¼ŒæŒ‰å¥å·åˆ†å‰²
            sentences = re.split(r'[ã€‚ï¼ï¼Ÿ]', description)
            result = []
            current_sentence = ""
            for sentence in sentences:
                if len(current_sentence + sentence) <= max_length:
                    current_sentence += sentence + "ã€‚"
                else:
                    if current_sentence:
                        result.append(current_sentence.strip())
                    current_sentence = sentence + "ã€‚"
            if current_sentence:
                result.append(current_sentence.strip())
            
            return result if result else [description]
        
        def get_unique_filename(filepath):
            """è·å–å”¯ä¸€æ–‡ä»¶å"""
            if not os.path.exists(filepath):
                return filepath
            
            base, ext = os.path.splitext(filepath)
            counter = 1
            while os.path.exists(f"{base}_{counter}{ext}"):
                counter += 1
            return f"{base}_{counter}{ext}"
        
        def clean_cell_value(value):
            """æ¸…ç†å•å…ƒæ ¼å€¼ï¼Œç§»é™¤éæ³•å­—ç¬¦"""
            if not value:
                return ""
            
            # è½¬æ¢ä¸ºå­—ç¬¦ä¸²
            value_str = str(value)
            
            # ç§»é™¤æ§åˆ¶å­—ç¬¦å’Œéæ³•å­—ç¬¦
            cleaned = ''
            for char in value_str:
                # åªä¿ç•™å¯æ‰“å°å­—ç¬¦å’Œå¸¸è§çš„ä¸­æ–‡å­—ç¬¦
                if char.isprintable() or '\u4e00' <= char <= '\u9fff':
                    cleaned += char
            
            # ç§»é™¤è°ƒè¯•ä¿¡æ¯ä¸­çš„ç‰¹æ®Šå­—ç¬¦
            cleaned = cleaned.replace('ğŸ”', '').replace('âœ…', '').replace('âš ï¸', '').replace('ğŸ“‹', '')
            
            # ç§»é™¤å¤šä½™çš„ç©ºç™½å­—ç¬¦
            cleaned = re.sub(r'\s+', ' ', cleaned.strip())
            
            # ç§»é™¤Excelä¸å…è®¸çš„ç‰¹æ®Šå­—ç¬¦
            cleaned = re.sub(r'[^\w\s\u4e00-\u9fff.,ï¼Œã€‚ï¼ï¼Ÿï¼š:()ï¼ˆï¼‰\-]', '', cleaned)
            
            return cleaned
        
        # æ¸…ç†æ•°æ®
        cleaned_data = self._clean_extracted_data(data)
        
        if not cleaned_data:
            logger.warning("æ¸…ç†åæ²¡æœ‰æ•°æ®éœ€è¦è¾“å‡º")
            return None
        
        # å¤„ç†é•¿æè¿°
        processed_data = []
        for item in cleaned_data:
            processed_item = item.copy()
            
            # å¤„ç†æ ‡ä¹¦æè¿°
            if processed_item.get('æ ‡ä¹¦æè¿°'):
                desc_parts = split_long_description(processed_item['æ ‡ä¹¦æè¿°'])
                for i, part in enumerate(desc_parts):
                    if i == 0:
                        processed_item['æ ‡ä¹¦æè¿°'] = part
                    else:
                        # åˆ›å»ºæ–°è¡Œ
                        new_item = processed_item.copy()
                        new_item['æ ‡ä¹¦æè¿°'] = part
                        processed_data.append(new_item)
                        processed_item = None
                        break
                if processed_item:
                    processed_data.append(processed_item)
            else:
                processed_data.append(processed_item)
        
        # åˆ›å»ºDataFrame
        df = pd.DataFrame(processed_data)
        
        # ç¡®ä¿åˆ—é¡ºåºæ­£ç¡®
        column_order = ['ä¸€çº§æ¨¡å—åç§°', 'äºŒçº§æ¨¡å—åç§°', 'ä¸‰çº§æ¨¡å—åç§°', 'æ ‡ä¹¦æè¿°', 'åˆåŒæè¿°', 'æ¥æºæ–‡ä»¶']
        for col in column_order:
            if col not in df.columns:
                df[col] = ''
        df = df[column_order]
        
        # æ¸…ç†å•å…ƒæ ¼å€¼
        for col in df.columns:
            df[col] = df[col].apply(clean_cell_value)
        
        # è·å–å”¯ä¸€æ–‡ä»¶å
        output_path = get_unique_filename(output_path)
        
        # å†™å…¥Excel
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='æå–ç»“æœ', index=False)
            worksheet = writer.sheets['æå–ç»“æœ']
            
            # è®¾ç½®åˆ—å®½
            column_widths = {'A': 15, 'B': 15, 'C': 15, 'D': 45, 'E': 45, 'F': 10}
            for col, width in column_widths.items():
                worksheet.column_dimensions[col].width = width
            
            # è®¾ç½®è¡¨å¤´æ ·å¼
            header_font = Font(bold=True, size=12)
            header_alignment = Alignment(horizontal='center', vertical='center')
            for col in range(1, len(column_order) + 1):
                cell = worksheet.cell(row=1, column=col)
                cell.font = header_font
                cell.alignment = header_alignment
            
            # è®¾ç½®æ•°æ®è¡Œæ ·å¼å’Œè¡Œé«˜
            data_alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
            for row in range(2, len(df) + 2):
                # è®¡ç®—æ¯è¡Œæœ€å¤§å­—ç¬¦æ•°
                max_chars = 0
                for col in range(1, len(column_order) + 1):
                    cell_value = str(worksheet.cell(row=row, column=col).value or '')
                    lines = cell_value.split('\n')
                    for line in lines:
                        line_chars = len(line)
                        if line_chars > 30:
                            needed_lines = (line_chars // 30) + 1
                            max_chars = max(max_chars, needed_lines * 30)
                else:
                            max_chars = max(max_chars, line_chars)
                
                # è®¡ç®—è¡Œé«˜
                estimated_lines = max(1, (max_chars // 30) + 1)
                row_height = max(20, estimated_lines * 18 + 10)
                worksheet.row_dimensions[row].height = row_height
                
                # è®¾ç½®å•å…ƒæ ¼å¯¹é½æ–¹å¼
                for col in range(1, len(column_order) + 1):
                    cell = worksheet.cell(row=row, column=col)
                    cell.alignment = data_alignment
            
            # è®¾ç½®è¾¹æ¡†
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            for row in worksheet.iter_rows(min_row=1, max_row=len(df) + 1, min_col=1, max_col=len(column_order)):
                for cell in row:
                    cell.border = thin_border
        
        return output_path

def main():
    """ä¸»å‡½æ•°ï¼ˆä¿ç•™ç”¨äºæµ‹è¯•ï¼‰"""
    extractor = PDFWordTableExtractor()
    print("PDFå’ŒWordè¡¨æ ¼æå–å·¥å…· - Webç‰ˆæœ¬")
    print("è¯·é€šè¿‡Streamlitåº”ç”¨ä½¿ç”¨æ­¤å·¥å…·")

if __name__ == "__main__":
    main()
# æµ‹è¯•ä»£ç 